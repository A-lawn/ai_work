"""DOCX 文档解析器"""
import logging
from typing import Dict, Any
from docx import Document
from app.parsers.base import DocumentParser

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """DOCX 文档解析器"""
    
    def parse(self, file_path: str) -> str:
        """
        解析 DOCX 文档并提取文本
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 保留标题层级信息
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        text_content.append(f"\n[{para.style.name}] {text}\n")
                    else:
                        text_content.append(text)
            
            # 提取表格文本
            for table_idx, table in enumerate(doc.tables, start=1):
                text_content.append(f"\n[Table {table_idx}]")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n".join(text_content)
            logger.info(
                f"成功解析 DOCX: {file_path}, "
                f"段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}"
            )
            
            return full_text
            
        except Exception as e:
            logger.error(f"解析 DOCX 失败: {file_path}, 错误: {str(e)}")
            raise Exception(f"DOCX 解析失败: {str(e)}")
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        获取 DOCX 文档元数据
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            文档元数据
        """
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            return {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
        except Exception as e:
            logger.warning(f"获取 DOCX 元数据失败: {str(e)}")
            return {}
    
    @property
    def supported_extensions(self) -> list:
        """支持的文件扩展名"""
        return [".docx", ".doc"]
